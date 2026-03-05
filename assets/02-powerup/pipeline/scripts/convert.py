"""
Stage 2: Document Conversion
=============================
Converts source files (md, csv, docx, pdf) to normalized markdown
with structural metadata, ready for the chunking step.

Replaces Docling MCP with lightweight Python libraries:
- python-docx for .docx
- pdfplumber for .pdf
- stdlib csv for .csv
- direct read for .md

Usage:
    python convert.py --input ./output/registered.json --output-dir ./output

Prerequisites:
    - pip install python-docx pdfplumber
"""

import argparse
import csv
import json
import re
import sys
from io import StringIO
from pathlib import Path


# ---------------------------------------------------------------------------
# Converters
# ---------------------------------------------------------------------------
def convert_md(file_path: str) -> tuple[str, dict]:
    """Read markdown as-is and extract heading structure."""
    with open(file_path, "r", encoding="utf-8") as f:
        content = f.read()

    headings = []
    for line in content.split("\n"):
        match = re.match(r"^(#{1,6})\s+(.+)$", line)
        if match:
            headings.append({
                "level": len(match.group(1)),
                "text": match.group(2).strip(),
            })

    return content, {"headings": headings}


def convert_csv(file_path: str) -> tuple[str, dict]:
    """Convert CSV to markdown table with column metadata."""
    with open(file_path, "r", encoding="utf-8") as f:
        reader = csv.reader(f)
        rows = list(reader)

    if not rows:
        return "", {"headings": [], "row_count": 0, "columns": []}

    columns = rows[0]
    data_rows = rows[1:]

    # Build markdown table
    lines = []
    lines.append("| " + " | ".join(columns) + " |")
    lines.append("| " + " | ".join(["---"] * len(columns)) + " |")
    for row in data_rows:
        # Pad row if shorter than header
        padded = row + [""] * (len(columns) - len(row))
        lines.append("| " + " | ".join(padded[:len(columns)]) + " |")

    markdown = "\n".join(lines)

    return markdown, {
        "headings": [],
        "row_count": len(data_rows),
        "columns": columns,
    }


def convert_docx(file_path: str) -> tuple[str, dict]:
    """Convert DOCX to markdown preserving heading hierarchy."""
    try:
        from docx import Document
    except ImportError:
        print("Error: python-docx not installed. Run: pip install python-docx")
        sys.exit(1)

    doc = Document(file_path)
    lines = []
    headings = []

    for para in doc.paragraphs:
        style_name = para.style.name.lower() if para.style else ""
        text = para.text.strip()
        if not text:
            continue

        if "heading 1" in style_name or style_name == "title":
            lines.append(f"# {text}")
            headings.append({"level": 1, "text": text})
        elif "heading 2" in style_name:
            lines.append(f"## {text}")
            headings.append({"level": 2, "text": text})
        elif "heading 3" in style_name:
            lines.append(f"### {text}")
            headings.append({"level": 3, "text": text})
        elif "heading 4" in style_name:
            lines.append(f"#### {text}")
            headings.append({"level": 4, "text": text})
        elif "list" in style_name:
            lines.append(f"- {text}")
        else:
            lines.append(text)

    # Handle tables
    for table in doc.tables:
        table_lines = []
        for i, row in enumerate(table.rows):
            cells = [cell.text.strip() for cell in row.cells]
            table_lines.append("| " + " | ".join(cells) + " |")
            if i == 0:
                table_lines.append("| " + " | ".join(["---"] * len(cells)) + " |")
        lines.append("\n".join(table_lines))

    markdown = "\n".join(lines)
    return markdown, {"headings": headings}


def convert_pdf(file_path: str) -> tuple[str, dict]:
    """Convert PDF to markdown with page boundaries."""
    try:
        import pdfplumber
    except ImportError:
        print("Error: pdfplumber not installed. Run: pip install pdfplumber")
        sys.exit(1)

    pages = []
    headings = []

    with pdfplumber.open(file_path) as pdf:
        page_count = len(pdf.pages)
        for i, page in enumerate(pdf.pages):
            text = page.extract_text() or ""
            if text.strip():
                pages.append(f"<!-- Page {i + 1} -->\n\n{text}")

            # Extract tables
            for table in page.extract_tables():
                if table and len(table) > 1:
                    table_lines = []
                    for j, row in enumerate(table):
                        cells = [str(c or "").strip() for c in row]
                        table_lines.append("| " + " | ".join(cells) + " |")
                        if j == 0:
                            table_lines.append(
                                "| " + " | ".join(["---"] * len(cells)) + " |"
                            )
                    pages.append("\n".join(table_lines))

    markdown = "\n\n".join(pages)

    # Try to detect headings from text patterns (lines in ALL CAPS or short bold lines)
    for line in markdown.split("\n"):
        stripped = line.strip()
        if not stripped or stripped.startswith("<!--") or stripped.startswith("|"):
            continue
        # Heuristic: short lines that look like section titles
        if len(stripped) < 80 and stripped == stripped.upper() and len(stripped) > 3:
            headings.append({"level": 1, "text": stripped.title()})

    return markdown, {"headings": headings, "page_count": page_count}


# ---------------------------------------------------------------------------
# Router
# ---------------------------------------------------------------------------
CONVERTERS = {
    "md": convert_md,
    "csv": convert_csv,
    "docx": convert_docx,
    "pdf": convert_pdf,
}


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------
def main():
    parser = argparse.ArgumentParser(description="RAG Pipeline -- Document Conversion")
    parser.add_argument(
        "--input",
        default="./output/registered.json",
        help="Path to registered.json from ingest step",
    )
    parser.add_argument(
        "--output-dir",
        default="./output",
        help="Path to write converted output (default: ./output)",
    )
    args = parser.parse_args()

    input_path = Path(args.input)
    if not input_path.exists():
        print(f"Error: Input file not found: {input_path}")
        print("Run ingest.py first.")
        sys.exit(1)

    with open(input_path, "r", encoding="utf-8") as f:
        files = json.load(f)

    output_dir = Path(args.output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    print("=" * 60)
    print("Stage 2: Document Conversion")
    print("=" * 60)

    converted = []
    failures = []

    for file_info in files:
        file_name = file_info["file_name"]
        file_type = file_info["file_type"]
        file_path = file_info["file_path"]

        converter = CONVERTERS.get(file_type)
        if not converter:
            print(f"  Skip (unsupported): {file_name} ({file_type})")
            continue

        print(f"  Converting: {file_name} ({file_type})...", end=" ")

        try:
            markdown, structure = converter(file_path)
            converted.append({
                "file_metadata_id": file_info["file_metadata_id"],
                "file_name": file_name,
                "file_type": file_type,
                "markdown": markdown,
                "structure": structure,
            })
            print(f"OK ({len(markdown)} chars)")
        except Exception as e:
            print(f"FAILED: {e}")
            failures.append({"file_name": file_name, "error": str(e)})

    # Write output
    output_file = output_dir / "converted.json"
    with open(output_file, "w", encoding="utf-8") as f:
        json.dump(converted, f, indent=2)

    print(f"\n{'=' * 60}")
    print(f"Done. {len(converted)} files converted, {len(failures)} failures.")
    if failures:
        for fail in failures:
            print(f"  FAILED: {fail['file_name']} -- {fail['error']}")
    print(f"Output: {output_file}")
    print(f"Next step: python chunking.py --input {output_file}")


if __name__ == "__main__":
    main()
